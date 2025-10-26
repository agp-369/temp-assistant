from src.plugin_interface import Plugin
from docx import Document
import os

class DocumentPlugin(Plugin):
    """
    A plugin to handle document creation.
    """
    def get_intent_map(self):
        return {
            "create_document": self.handle
        }

    def can_handle(self, command):
        return False # This plugin is intent-based

    def handle(self, command, assistant):
        intent, args = command

        # We need to store the content from the previous step (summarization)
        # This will require enhancing the context passed between plan steps.
        # For now, we'll use a placeholder.
        content = assistant.last_summary or "No content was provided for the document."

        # The argument is the topic
        topic = args.replace("about ", "")
        filename = f"{topic.replace(' ', '_')}_report.docx"

        try:
            document = Document()
            document.add_heading(topic.title(), 0)
            document.add_paragraph(content)

            # Save to the user's Documents folder
            save_path = os.path.join(os.path.expanduser("~"), "Documents", filename)
            document.save(save_path)

            assistant.speak(f"I have created the document and saved it as '{filename}' in your Documents folder.")

        except Exception as e:
            assistant.speak(f"I'm sorry, I encountered an error while creating the document: {e}", is_error=True)
